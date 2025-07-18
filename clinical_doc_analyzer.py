# Streamlit app: Inpatient CDI Analyzer
import streamlit as st
import fitz  # PyMuPDF
import docx
import tempfile
import os
import pandas as pd
from google.generativeai import GenerativeModel
import google.generativeai as genai
import json
import time
from io import BytesIO

# --- Load Gemini API Key ---
# This assumes the API key is set in Streamlit secrets.toml
# For local testing, you might hardcode it or use an environment variable.
if 'gemini' in st.secrets:
    genai.configure(api_key=st.secrets['gemini']['api_key'])
else:
    st.error("❌ Gemini API key not configured. Please set it in .streamlit/secrets.toml")
    st.stop()

model = genai.GenerativeModel('gemini-pro')

# --- Helper Functions ---

def extract_text_from_pdf(file):
    """Extracts text from a PDF file using PyMuPDF."""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        text = "\n".join(page.get_text() for page in doc)
    return text

def extract_text_from_docx(file):
    """Extracts text from a DOCX file using python-docx."""
    d = docx.Document(file)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())

def split_sections(text):
    """
    Splits the full clinical document text into logical sections based on common headers.
    Headers are case-insensitive.
    """
    # Common clinical document headers to identify sections
    headers = [
        "history of present illness", "progress note", "operative report",
        "discharge summary", "consultation", "physical exam",
        "assessment and plan", "emergency department record",
        "diagnostic test results", "physician orders",
        "anaesthesiology record", "medication administration record",
        "vital signs flowsheet", "problem list", "ventilator flowsheet",
        "wound care notes"
    ]
    lines = text.lower().splitlines()
    sections = {}
    current_section = "unknown_introduction" # Default section for text before any recognized header
    buffer = []

    for line in lines:
        found_header = False
        for h in headers:
            # Check if the line contains a header, allowing for some flexibility (e.g., "History of Present Illness:")
            if h in line and len(line) < 100: # Heuristic: header lines are usually short
                if buffer:
                    sections[current_section.strip()] = "\n".join(buffer).strip()
                    buffer = []
                current_section = line.strip()
                found_header = True
                break
        if not found_header:
            buffer.append(line)

    # Add any remaining text to the last section or 'unknown_introduction'
    if buffer:
        sections[current_section.strip()] = "\n".join(buffer).strip()
    return sections


def categorize_section_type(text):
    """
    Uses Gemini to categorize the type of clinical documentation section.
    Provides a document type and a confidence score.
    """
    snippet = text[:1500] # Use a larger snippet for better context
    prompt = f"""
    You are an expert in clinical documentation. Analyze the following text snippet from a patient's chart.
    What type of clinical documentation is this?
    Choose one from the most common types: History and Physical (H&P), Progress Note, Operative Report, Discharge Summary, Consultation Report, Emergency Department Record, Diagnostic Test Results, Physician Orders, Anesthesiology Record, Medication Administration Record, Vital Signs Flowsheet, Problem List, Ventilator Flowsheet, Wound Care Notes, or Other/Unknown.
    Provide a confidence percentage for your categorization.

    --- Text Snippet ---
    {snippet}
    --- End Snippet ---
    Respond in JSON format like this: {{"Document Type": "str", "Confidence": int}}
    """
    try:
        response = model.generate_content(prompt)
        # Attempt to parse JSON, handle cases where Gemini might return extra text
        json_str = response.text.strip().replace("```json\n", "").replace("\n```", "")
        return json.loads(json_str)
    except Exception as e:
        st.warning(f"Error categorizing section type: {e}. Gemini response: {response.text}")
        return {"Document Type": "Unknown", "Confidence": 0}


def generate_gemini_analysis(text, retries=3):
    """
    Uses Gemini to perform a detailed CDI analysis on the provided clinical text.
    Evaluates against 5 CDI criteria, provides scores, explanations, proposed actions,
    and draft queries.
    """
    # Detailed definitions of CDI criteria, derived from the uploaded RTF documents,
    # embedded directly into the prompt to guide Gemini's reasoning.
    cdi_criteria_definitions = """
    Here are the definitions for the 5 CDI-relevant criteria you must assess, drawing from the principles of high-quality clinical documentation:

    1.  **Clinical Reliability (Score 0-10):**
        * **Definition:** The content of the record is trustworthy, safe, and internally consistent, supporting the treatments provided. Documentation should logically align with the patient's condition and the interventions performed.
        * **Example of Deficiency:** Treatment provided without documentation of the condition being treated (e.g., Lasix given, but no CHF documented; KCI administered, no hypokalemia documented). Or, a diagnosis that does not appear reliable based on the treatment (e.g., blood transfusion for a bleeding gastric ulcer without acute blood loss anemia).
        * **Goal:** Ensure diagnoses and treatments are mutually supportive and clinically sound.

    2.  **Clinical Precision (Score 0-10):**
        * **Definition:** The record is accurate, exact, and strictly defined. Documentation should include specific details, etiology, and highest level of specificity where clinically appropriate. Avoid vague or general terms.
        * **Example of Deficiency:** No specific diagnosis documented when a more specific diagnosis appears to be supported (e.g., "anemia" vs. "acute or chronic blood loss anemia"; "pneumonia" vs. "aspiration pneumonia").
        * **Goal:** Drive documentation to the highest level of detail and specificity for accurate coding and reflection of patient acuity.

    3.  **Documentation Completeness (Score 0-10):**
        * **Definition:** The record has the maximum content and is thorough. All concerns are fully addressed, including appropriate authentication (signature, date). Abnormal test results should have documented clinical significance or be noted as clinically insignificant. Reasons for tests/treatments should be documented.
        * **Example of Deficiency:** Abnormal test results without documentation for clinical significance (e.g., low sodium, magnesium, potassium without corresponding diagnoses or notation of clinical insignificance). Lack of documented reason for ordered tests or treatments.
        * **Goal:** Ensure all relevant clinical information is present, authenticated, and fully explained.

    4.  **Clinical Consistency & Clarity (Score 0-10):**
        * **Definition:** The record is not contradictory, vague, or ambiguous. Documentation should be coherent across different entries and authors. If conflicting information exists, it should be resolved or clarified. Symptoms should ideally have identified etiologies; if not, "etiology undetermined" should be documented.
        * **Example of Deficiency:** Disagreement between two or more treating physicians without obvious resolution (e.g., primary care documents TIA, neurologist documents CVA, and attending physician doesn't clarify). Vague or ambiguous documentation, especially for symptom principal diagnoses (e.g., "chest pain" without further insight; "syncope" without etiology).
        * **Goal:** Achieve a unified, unambiguous clinical picture that accurately reflects the patient's condition and care.

    5.  **Timeliness (Flag Only - "Present", "Absent", or "N/A"):**
        * **Definition:** Documentation is prepared by the provider at the point it is needed for patient care, adhering to facility guidelines. This includes timely progress notes, discharge summaries, and documentation of Present on Admission (POA) indicators.
        * **Flagging:** You will only flag if there are *clear indicators* of timeliness issues (e.g., explicit mentions of late documentation, missing dates where expected). If no such indicators are present, assume "Absent" or "N/A" if the document type doesn't typically have strict timeliness requirements visible in the text.
        * **Goal:** Ensure documentation is current and available when needed for care, coding, and compliance.
    """

    prompt = f"""
    You are a Senior Clinical Documentation Integrity (CDI) Specialist. Your task is to systematically assess the following clinical document section using the provided theory of high-quality documentation.
    Apply your expertise to validate the documentation and identify any disturbances or areas for improvement.

    {cdi_criteria_definitions}

    For each of the first four criteria (Clinical Reliability, Clinical Precision, Documentation Completeness, Clinical Consistency & Clarity):
    - Assign a score (0-10), where 10 is perfect and 0 is completely deficient.
    - Explain "Why" this score was given, referencing specific issues or strengths in the document.
    - Suggest a "Proposed CDI Action" (e.g., "Educate provider on specificity," "Query for clinical validation," "Review for conflicting documentation").
    - If the score is less than 10, draft a clear and compliant "Draft Query" that a CDI specialist would send to the physician to clarify or improve the documentation. The query should be open-ended and non-leading.

    For the "Timeliness" criterion:
    - Flag it as "Present" if there are clear indications of timeliness issues (e.g., explicit mentions of late entries, missing dates where expected).
    - Flag it as "Absent" if no such issues are evident from the text.
    - Flag it as "N/A" if the document type or content doesn't provide enough information to assess timeliness.

    Return your output in structured JSON like this. Ensure all fields are present, even if empty or "N/A".

    {{
      "Scorecard": {{
        "Clinical Reliability": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Clinical Precision": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Documentation Completeness": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Clinical Consistency & Clarity": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}}
      }},
      "Timeliness Flag": "Present" | "Absent" | "N/A",
      "Final Summary": "Overall assessment of documentation quality for this section, highlighting key strengths and areas for CDI focus."
    }}

    --- Document Section Start ---
    {text}
    --- Document Section End ---
    """
    for attempt in range(retries):
        try:
            response = model.generate_content(prompt)
            # Clean the response to ensure it's valid JSON
            json_str = response.text.strip().replace("```json\n", "").replace("\n```", "")
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            st.warning(f"Attempt {attempt+1}: JSON decoding error: {e}. Retrying...")
            time.sleep(2) # Wait before retrying
        except Exception as e:
            st.warning(f"Attempt {attempt+1}: Gemini generation error: {e}. Retrying...")
            time.sleep(2) # Wait before retrying
    return {"error": "Gemini failed after retries or returned invalid JSON.", "raw_response": response.text if 'response' in locals() else "No response"}


def generate_excel(results):
    """
    Generates a Pandas DataFrame from the analysis results, suitable for Excel export.
    Includes all CDI criteria, scores, explanations, actions, and queries.
    """
    df_rows = []
    for section, result in results.items():
        scorecard = result.get("Scorecard", {})
        doc_type = result.get("Doc Type", "")
        confidence = result.get("Confidence", "")
        timeliness_flag = result.get("Timeliness Flag", "N/A")
        final_summary = result.get("Final Summary", "")

        # Add a row for the overall section summary and timeliness flag
        df_rows.append({
            "Section": section.title(),
            "AI Doc Type": doc_type,
            "Confidence %": confidence,
            "Criterion": "Overall Section Summary",
            "Score": "-", # No score for summary row
            "Why": final_summary,
            "Proposed CDI Action": "",
            "Draft Query": "",
            "Timeliness Flag": timeliness_flag # Display timeliness here
        })

        # Add rows for each detailed scorecard criterion
        for criterion, detail in scorecard.items():
            score = detail.get("Score", "")
            # Only include Draft Query if score is less than 10
            draft_query = detail.get("Draft Query", "") if int(str(score).replace('-', '0')) < 10 else ""
            df_rows.append({
                "Section": section.title(),
                "AI Doc Type": doc_type,
                "Confidence %": confidence,
                "Criterion": criterion,
                "Score": score,
                "Why": detail.get("Why", ""),
                "Proposed CDI Action": detail.get("Proposed CDI Action", ""),
                "Draft Query": draft_query,
                "Timeliness Flag": timeliness_flag # Repeat for consistency in rows
            })
    return pd.DataFrame(df_rows)

def highlight_score(val):
    """
    Applies conditional formatting to the 'Score' column in the DataFrame.
    Green for good scores, yellow for moderate, red for low.
    """
    try:
        score = int(val)
        if score >= 9:
            return 'background-color: #d9ead3'  # Light Green
        elif score >= 7:
            return 'background-color: #fff2cc'  # Light Yellow
        else:
            return 'background-color: #fce5cd'  # Light Orange
    except ValueError: # Handle non-integer values like '-'
        return ''

# --- Streamlit UI ---
st.set_page_config(page_title="Inpatient CDI Analyzer", layout="wide")
st.title("🩺 Inpatient CDI Analyzer: AI-Powered Clinical Documentation Review")
st.markdown("Upload a full inpatient PDF or Word document to get a detailed CDI analysis, including scores, proposed actions, and draft queries.")

uploaded_file = st.file_uploader("📤 Upload Full Inpatient Chart (.pdf or .docx)", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("Processing document..."):
        if uploaded_file.name.endswith(".pdf"):
            full_text = extract_text_from_pdf(uploaded_file)
        else:
            full_text = extract_text_from_docx(uploaded_file)

    st.success("✅ Document Uploaded and Processed!")

    sections = split_sections(full_text)
    st.subheader("📑 Detected Sections")
    st.write(f"Found {len(sections)} sections in the document.")

    all_results = {}

    # Button to analyze all sections
    if st.button("🧠 Analyze All Sections (This may take a few minutes for large documents)"):
        progress_bar = st.progress(0)
        total_sections = len(sections)
        for i, (section_name, section_text) in enumerate(sections.items()):
            st.info(f"Analyzing section: **{section_name.title()}** ({i+1}/{total_sections})...")
            with st.spinner(f"Analyzing: {section_name.title()}..."):
                doc_info = categorize_section_type(section_text)
                analysis_result = generate_gemini_analysis(section_text)

                if 'error' not in analysis_result and analysis_result.get("Scorecard"):
                    analysis_result["Doc Type"] = doc_info.get("Document Type")
                    analysis_result["Confidence"] = doc_info.get("Confidence")
                    all_results[section_name] = analysis_result
                    st.success(f"✅ Analysis complete for **{section_name.title()}**.")
                else:
                    st.error(f"❌ Failed to analyze section: **{section_name.title()}**. Error: {analysis_result.get('error', 'Unknown')}")
            progress_bar.progress((i + 1) / total_sections)
        st.success("🎉 All sections analyzed!")

        # Display results in a DataFrame
        if all_results:
            st.subheader("📊 Comprehensive CDI Analysis Results")
            df = generate_excel(all_results)

            if 'Score' in df.columns:
                # Apply styling only to rows that are actual criterion scores, not summary rows
                styled_df = df.style.applymap(highlight_score, subset=['Score'])
                st.dataframe(styled_df, use_container_width=True)
            else:
                st.warning("⚠️ No valid 'Score' column found in results. Showing raw output.")
                st.dataframe(df, use_container_width=True)

            # Download buttons
            col1, col2 = st.columns(2)
            with col1:
                with BytesIO() as buffer:
                    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                        df.to_excel(writer, index=False)
                    st.download_button("📥 Download Excel Report", buffer.getvalue(), file_name="CDI_Analysis_Report.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col2:
                txt_summary = ""
                for sec, res in all_results.items():
                    txt_summary += f"\n\n=== {sec.title()} (AI Doc Type: {res.get('Doc Type', '')}, Confidence: {res.get('Confidence', '')}%) ===\n"
                    txt_summary += f"Overall Summary: {res.get('Final Summary', 'N/A')}\n"
                    txt_summary += f"Timeliness Flag: {res.get('Timeliness Flag', 'N/A')}\n"
                    if 'error' in res:
                        txt_summary += "Error during analysis.\n"
                    else:
                        for k, v in res.get("Scorecard", {}).items():
                            score = v.get('Score', '-')
                            why = v.get('Why', '')
                            action = v.get('Proposed CDI Action', '')
                            query = v.get('Draft Query', '')

                            txt_summary += f"\n--- {k} (Score: {score}/10) ---\n"
                            txt_summary += f"Why: {why}\n"
                            txt_summary += f"Proposed CDI Action: {action}\n"
                            if query: # Only add query if it's not empty
                                txt_summary += f"→ Draft Query: {query}\n"
                st.download_button("📄 Download TXT Report", txt_summary, file_name="CDI_Queries_Report.txt", mime="text/plain")
        else:
            st.info("No analysis results to display yet. Upload a document and click 'Analyze All Sections'.")

    st.markdown("---")
    st.subheader("🔍 Analyze Individual Sections")
    st.info("You can also expand a section below and analyze it individually.")

    # Display individual sections with an analyze button for each
    for i, (section_name, section_text) in enumerate(sections.items()):
        with st.expander(f"📄 Section {i+1}: {section_name.title()}"):
            st.text_area(f"Content of '{section_name.title()}' (first 2000 chars)", section_text[:2000], height=200, key=f"text_area_{i}")
            if st.button(f"🔍 Analyze '{section_name.title()}' Individually", key=f"analyze_single_{i}"):
                with st.spinner(f"Analyzing '{section_name.title()}' with Gemini..."):
                    doc_info = categorize_section_type(section_text)
                    result = generate_gemini_analysis(section_text)
                    result["Doc Type"] = doc_info.get("Document Type")
                    result["Confidence"] = doc_info.get("Confidence")

                    if 'error' not in result:
                        st.success(f"✅ Analysis for '{section_name.title()}' complete!")
                        st.json(result) # Display raw JSON for single section analysis
                    else:
                        st.error(f"❌ Failed to analyze '{section_name.title()}'. Error: {result.get('error', 'Unknown')}")
                        if 'raw_response' in result:
                            st.code(result['raw_response'], language='json') # Show raw response for debugging
# Streamlit app: Inpatient CDI Analyzer
import streamlit as st
import fitz  # PyMuPDF
import docx
import tempfile
import os
import pandas as pd
from google.generativeai import GenerativeModel
import google.generativeai as genai
import json
import time
from io import BytesIO

# --- Load Gemini API Key ---
# This assumes the API key is set in Streamlit secrets.toml
# For local testing, you might hardcode it or use an environment variable.
if 'gemini' in st.secrets:
    genai.configure(api_key=st.secrets['gemini']['api_key'])
else:
    st.error("❌ Gemini API key not configured. Please set it in .streamlit/secrets.toml")
    st.stop()

model = genai.GenerativeModel('gemini-pro')

# --- Helper Functions ---

def extract_text_from_pdf(file):
    """Extracts text from a PDF file using PyMuPDF."""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        text = "\n".join(page.get_text() for page in doc)
    return text

def extract_text_from_docx(file):
    """Extracts text from a DOCX file using python-docx."""
    d = docx.Document(file)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())

def split_sections(text):
    """
    Splits the full clinical document text into logical sections based on common headers.
    Headers are case-insensitive.
    """
    # Common clinical document headers to identify sections
    headers = [
        "history of present illness", "progress note", "operative report",
        "discharge summary", "consultation", "physical exam",
        "assessment and plan", "emergency department record",
        "diagnostic test results", "physician orders",
        "anaesthesiology record", "medication administration record",
        "vital signs flowsheet", "problem list", "ventilator flowsheet",
        "wound care notes"
    ]
    lines = text.lower().splitlines()
    sections = {}
    current_section = "unknown_introduction" # Default section for text before any recognized header
    buffer = []

    for line in lines:
        found_header = False
        for h in headers:
            # Check if the line contains a header, allowing for some flexibility (e.g., "History of Present Illness:")
            if h in line and len(line) < 100: # Heuristic: header lines are usually short
                if buffer:
                    sections[current_section.strip()] = "\n".join(buffer).strip()
                    buffer = []
                current_section = line.strip()
                found_header = True
                break
        if not found_header:
            buffer.append(line)

    # Add any remaining text to the last section or 'unknown_introduction'
    if buffer:
        sections[current_section.strip()] = "\n".join(buffer).strip()
    return sections


def categorize_section_type(text):
    """
    Uses Gemini to categorize the type of clinical documentation section.
    Provides a document type and a confidence score.
    """
    snippet = text[:1500] # Use a larger snippet for better context
    prompt = f"""
    You are an expert in clinical documentation. Analyze the following text snippet from a patient's chart.
    What type of clinical documentation is this?
    Choose one from the most common types: History and Physical (H&P), Progress Note, Operative Report, Discharge Summary, Consultation Report, Emergency Department Record, Diagnostic Test Results, Physician Orders, Anesthesiology Record, Medication Administration Record, Vital Signs Flowsheet, Problem List, Ventilator Flowsheet, Wound Care Notes, or Other/Unknown.
    Provide a confidence percentage for your categorization.

    --- Text Snippet ---
    {snippet}
    --- End Snippet ---
    Respond in JSON format like this: {{"Document Type": "str", "Confidence": int}}
    """
    try:
        response = model.generate_content(prompt)
        # Attempt to parse JSON, handle cases where Gemini might return extra text
        json_str = response.text.strip().replace("```json\n", "").replace("\n```", "")
        return json.loads(json_str)
    except Exception as e:
        st.warning(f"Error categorizing section type: {e}. Gemini response: {response.text}")
        return {"Document Type": "Unknown", "Confidence": 0}


def generate_gemini_analysis(text, retries=3):
    """
    Uses Gemini to perform a detailed CDI analysis on the provided clinical text.
    Evaluates against 5 CDI criteria, provides scores, explanations, proposed actions,
    and draft queries.
    """
    # Detailed definitions of CDI criteria, derived from the uploaded RTF documents,
    # embedded directly into the prompt to guide Gemini's reasoning.
    cdi_criteria_definitions = """
    Here are the definitions for the 5 CDI-relevant criteria you must assess, drawing from the principles of high-quality clinical documentation:

    1.  **Clinical Reliability (Score 0-10):**
        * **Definition:** The content of the record is trustworthy, safe, and internally consistent, supporting the treatments provided. Documentation should logically align with the patient's condition and the interventions performed.
        * **Example of Deficiency:** Treatment provided without documentation of the condition being treated (e.g., Lasix given, but no CHF documented; KCI administered, no hypokalemia documented). Or, a diagnosis that does not appear reliable based on the treatment (e.g., blood transfusion for a bleeding gastric ulcer without acute blood loss anemia).
        * **Goal:** Ensure diagnoses and treatments are mutually supportive and clinically sound.

    2.  **Clinical Precision (Score 0-10):**
        * **Definition:** The record is accurate, exact, and strictly defined. Documentation should include specific details, etiology, and highest level of specificity where clinically appropriate. Avoid vague or general terms.
        * **Example of Deficiency:** No specific diagnosis documented when a more specific diagnosis appears to be supported (e.g., "anemia" vs. "acute or chronic blood loss anemia"; "pneumonia" vs. "aspiration pneumonia").
        * **Goal:** Drive documentation to the highest level of detail and specificity for accurate coding and reflection of patient acuity.

    3.  **Documentation Completeness (Score 0-10):**
        * **Definition:** The record has the maximum content and is thorough. All concerns are fully addressed, including appropriate authentication (signature, date). Abnormal test results should have documented clinical significance or be noted as clinically insignificant. Reasons for tests/treatments should be documented.
        * **Example of Deficiency:** Abnormal test results without documentation for clinical significance (e.g., low sodium, magnesium, potassium without corresponding diagnoses or notation of clinical insignificance). Lack of documented reason for ordered tests or treatments.
        * **Goal:** Ensure all relevant clinical information is present, authenticated, and fully explained.

    4.  **Clinical Consistency & Clarity (Score 0-10):**
        * **Definition:** The record is not contradictory, vague, or ambiguous. Documentation should be coherent across different entries and authors. If conflicting information exists, it should be resolved or clarified. Symptoms should ideally have identified etiologies; if not, "etiology undetermined" should be documented.
        * **Example of Deficiency:** Disagreement between two or more treating physicians without obvious resolution (e.g., primary care documents TIA, neurologist documents CVA, and attending physician doesn't clarify). Vague or ambiguous documentation, especially for symptom principal diagnoses (e.g., "chest pain" without further insight; "syncope" without etiology).
        * **Goal:** Achieve a unified, unambiguous clinical picture that accurately reflects the patient's condition and care.

    5.  **Timeliness (Flag Only - "Present", "Absent", or "N/A"):**
        * **Definition:** Documentation is prepared by the provider at the point it is needed for patient care, adhering to facility guidelines. This includes timely progress notes, discharge summaries, and documentation of Present on Admission (POA) indicators.
        * **Flagging:** You will only flag if there are *clear indicators* of timeliness issues (e.g., explicit mentions of late documentation, missing dates where expected). If no such indicators are present, assume "Absent" or "N/A" if the document type doesn't typically have strict timeliness requirements visible in the text.
        * **Goal:** Ensure documentation is current and available when needed for care, coding, and compliance.
    """

    prompt = f"""
    You are a Senior Clinical Documentation Integrity (CDI) Specialist. Your task is to systematically assess the following clinical document section using the provided theory of high-quality documentation.
    Apply your expertise to validate the documentation and identify any disturbances or areas for improvement.

    {cdi_criteria_definitions}

    For each of the first four criteria (Clinical Reliability, Clinical Precision, Documentation Completeness, Clinical Consistency & Clarity):
    - Assign a score (0-10), where 10 is perfect and 0 is completely deficient.
    - Explain "Why" this score was given, referencing specific issues or strengths in the document.
    - Suggest a "Proposed CDI Action" (e.g., "Educate provider on specificity," "Query for clinical validation," "Review for conflicting documentation").
    - If the score is less than 10, draft a clear and compliant "Draft Query" that a CDI specialist would send to the physician to clarify or improve the documentation. The query should be open-ended and non-leading.

    For the "Timeliness" criterion:
    - Flag it as "Present" if there are clear indications of timeliness issues (e.g., explicit mentions of late entries, missing dates where expected).
    - Flag it as "Absent" if no such issues are evident from the text.
    - Flag it as "N/A" if the document type or content doesn't provide enough information to assess timeliness.

    Return your output in structured JSON like this. Ensure all fields are present, even if empty or "N/A".

    {{
      "Scorecard": {{
        "Clinical Reliability": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Clinical Precision": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Documentation Completeness": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Clinical Consistency & Clarity": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}}
      }},
      "Timeliness Flag": "Present" | "Absent" | "N/A",
      "Final Summary": "Overall assessment of documentation quality for this section, highlighting key strengths and areas for CDI focus."
    }}

    --- Document Section Start ---
    {text}
    --- Document Section End ---
    """
    for attempt in range(retries):
        try:
            response = model.generate_content(prompt)
            # Clean the response to ensure it's valid JSON
            json_str = response.text.strip().replace("```json\n", "").replace("\n```", "")
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            st.warning(f"Attempt {attempt+1}: JSON decoding error: {e}. Retrying...")
            time.sleep(2) # Wait before retrying
        except Exception as e:
            st.warning(f"Attempt {attempt+1}: Gemini generation error: {e}. Retrying...")
            time.sleep(2) # Wait before retrying
    return {"error": "Gemini failed after retries or returned invalid JSON.", "raw_response": response.text if 'response' in locals() else "No response"}


def generate_excel(results):
    """
    Generates a Pandas DataFrame from the analysis results, suitable for Excel export.
    Includes all CDI criteria, scores, explanations, actions, and queries.
    """
    df_rows = []
    for section, result in results.items():
        scorecard = result.get("Scorecard", {})
        doc_type = result.get("Doc Type", "")
        confidence = result.get("Confidence", "")
        timeliness_flag = result.get("Timeliness Flag", "N/A")
        final_summary = result.get("Final Summary", "")

        # Add a row for the overall section summary and timeliness flag
        df_rows.append({
            "Section": section.title(),
            "AI Doc Type": doc_type,
            "Confidence %": confidence,
            "Criterion": "Overall Section Summary",
            "Score": "-", # No score for summary row
            "Why": final_summary,
            "Proposed CDI Action": "",
            "Draft Query": "",
            "Timeliness Flag": timeliness_flag # Display timeliness here
        })

        # Add rows for each detailed scorecard criterion
        for criterion, detail in scorecard.items():
            score = detail.get("Score", "")
            # Only include Draft Query if score is less than 10
            draft_query = detail.get("Draft Query", "") if int(str(score).replace('-', '0')) < 10 else ""
            df_rows.append({
                "Section": section.title(),
                "AI Doc Type": doc_type,
                "Confidence %": confidence,
                "Criterion": criterion,
                "Score": score,
                "Why": detail.get("Why", ""),
                "Proposed CDI Action": detail.get("Proposed CDI Action", ""),
                "Draft Query": draft_query,
                "Timeliness Flag": timeliness_flag # Repeat for consistency in rows
            })
    return pd.DataFrame(df_rows)

def highlight_score(val):
    """
    Applies conditional formatting to the 'Score' column in the DataFrame.
    Green for good scores, yellow for moderate, red for low.
    """
    try:
        score = int(val)
        if score >= 9:
            return 'background-color: #d9ead3'  # Light Green
        elif score >= 7:
            return 'background-color: #fff2cc'  # Light Yellow
        else:
            return 'background-color: #fce5cd'  # Light Orange
    except ValueError: # Handle non-integer values like '-'
        return ''

# --- Streamlit UI ---
st.set_page_config(page_title="Inpatient CDI Analyzer", layout="wide")
st.title("🩺 Inpatient CDI Analyzer: AI-Powered Clinical Documentation Review")
st.markdown("Upload a full inpatient PDF or Word document to get a detailed CDI analysis, including scores, proposed actions, and draft queries.")

uploaded_file = st.file_uploader("📤 Upload Full Inpatient Chart (.pdf or .docx)", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("Processing document..."):
        if uploaded_file.name.endswith(".pdf"):
            full_text = extract_text_from_pdf(uploaded_file)
        else:
            full_text = extract_text_from_docx(uploaded_file)

    st.success("✅ Document Uploaded and Processed!")

    sections = split_sections(full_text)
    st.subheader("📑 Detected Sections")
    st.write(f"Found {len(sections)} sections in the document.")

    all_results = {}

    # Button to analyze all sections
    if st.button("🧠 Analyze All Sections (This may take a few minutes for large documents)"):
        progress_bar = st.progress(0)
        total_sections = len(sections)
        for i, (section_name, section_text) in enumerate(sections.items()):
            st.info(f"Analyzing section: **{section_name.title()}** ({i+1}/{total_sections})...")
            with st.spinner(f"Analyzing: {section_name.title()}..."):
                doc_info = categorize_section_type(section_text)
                analysis_result = generate_gemini_analysis(section_text)

                if 'error' not in analysis_result and analysis_result.get("Scorecard"):
                    analysis_result["Doc Type"] = doc_info.get("Document Type")
                    analysis_result["Confidence"] = doc_info.get("Confidence")
                    all_results[section_name] = analysis_result
                    st.success(f"✅ Analysis complete for **{section_name.title()}**.")
                else:
                    st.error(f"❌ Failed to analyze section: **{section_name.title()}**. Error: {analysis_result.get('error', 'Unknown')}")
            progress_bar.progress((i + 1) / total_sections)
        st.success("🎉 All sections analyzed!")

        # Display results in a DataFrame
        if all_results:
            st.subheader("📊 Comprehensive CDI Analysis Results")
            df = generate_excel(all_results)

            if 'Score' in df.columns:
                # Apply styling only to rows that are actual criterion scores, not summary rows
                styled_df = df.style.applymap(highlight_score, subset=['Score'])
                st.dataframe(styled_df, use_container_width=True)
            else:
                st.warning("⚠️ No valid 'Score' column found in results. Showing raw output.")
                st.dataframe(df, use_container_width=True)

            # Download buttons
            col1, col2 = st.columns(2)
            with col1:
                with BytesIO() as buffer:
                    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                        df.to_excel(writer, index=False)
                    st.download_button("📥 Download Excel Report", buffer.getvalue(), file_name="CDI_Analysis_Report.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col2:
                txt_summary = ""
                for sec, res in all_results.items():
                    txt_summary += f"\n\n=== {sec.title()} (AI Doc Type: {res.get('Doc Type', '')}, Confidence: {res.get('Confidence', '')}%) ===\n"
                    txt_summary += f"Overall Summary: {res.get('Final Summary', 'N/A')}\n"
                    txt_summary += f"Timeliness Flag: {res.get('Timeliness Flag', 'N/A')}\n"
                    if 'error' in res:
                        txt_summary += "Error during analysis.\n"
                    else:
                        for k, v in res.get("Scorecard", {}).items():
                            score = v.get('Score', '-')
                            why = v.get('Why', '')
                            action = v.get('Proposed CDI Action', '')
                            query = v.get('Draft Query', '')

                            txt_summary += f"\n--- {k} (Score: {score}/10) ---\n"
                            txt_summary += f"Why: {why}\n"
                            txt_summary += f"Proposed CDI Action: {action}\n"
                            if query: # Only add query if it's not empty
                                txt_summary += f"→ Draft Query: {query}\n"
                st.download_button("📄 Download TXT Report", txt_summary, file_name="CDI_Queries_Report.txt", mime="text/plain")
        else:
            st.info("No analysis results to display yet. Upload a document and click 'Analyze All Sections'.")

    st.markdown("---")
    st.subheader("🔍 Analyze Individual Sections")
    st.info("You can also expand a section below and analyze it individually.")

    # Display individual sections with an analyze button for each
    for i, (section_name, section_text) in enumerate(sections.items()):
        with st.expander(f"📄 Section {i+1}: {section_name.title()}"):
            st.text_area(f"Content of '{section_name.title()}' (first 2000 chars)", section_text[:2000], height=200, key=f"text_area_{i}")
            if st.button(f"🔍 Analyze '{section_name.title()}' Individually", key=f"analyze_single_{i}"):
                with st.spinner(f"Analyzing '{section_name.title()}' with Gemini..."):
                    doc_info = categorize_section_type(section_text)
                    result = generate_gemini_analysis(section_text)
                    result["Doc Type"] = doc_info.get("Document Type")
                    result["Confidence"] = doc_info.get("Confidence")

                    if 'error' not in result:
                        st.success(f"✅ Analysis for '{section_name.title()}' complete!")
                        st.json(result) # Display raw JSON for single section analysis
                    else:
                        st.error(f"❌ Failed to analyze '{section_name.title()}'. Error: {result.get('error', 'Unknown')}")
                        if 'raw_response' in result:
                            st.code(result['raw_response'], language='json') # Show raw response for debugging
# Streamlit app: Inpatient CDI Analyzer
import streamlit as st
import fitz  # PyMuPDF
import docx
import tempfile
import os
import pandas as pd
from google.generativeai import GenerativeModel
import google.generativeai as genai
import json
import time
from io import BytesIO

# --- Load Gemini API Key ---
# This assumes the API key is set in Streamlit secrets.toml
# For local testing, you might hardcode it or use an environment variable.
if 'gemini' in st.secrets:
    genai.configure(api_key=st.secrets['gemini']['api_key'])
else:
    st.error("❌ Gemini API key not configured. Please set it in .streamlit/secrets.toml")
    st.stop()

model = genai.GenerativeModel('gemini-pro')

# --- Helper Functions ---

def extract_text_from_pdf(file):
    """Extracts text from a PDF file using PyMuPDF."""
    with fitz.open(stream=file.read(), filetype="pdf") as doc:
        text = "\n".join(page.get_text() for page in doc)
    return text

def extract_text_from_docx(file):
    """Extracts text from a DOCX file using python-docx."""
    d = docx.Document(file)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())

def split_sections(text):
    """
    Splits the full clinical document text into logical sections based on common headers.
    Headers are case-insensitive.
    """
    # Common clinical document headers to identify sections
    headers = [
        "history of present illness", "progress note", "operative report",
        "discharge summary", "consultation", "physical exam",
        "assessment and plan", "emergency department record",
        "diagnostic test results", "physician orders",
        "anaesthesiology record", "medication administration record",
        "vital signs flowsheet", "problem list", "ventilator flowsheet",
        "wound care notes"
    ]
    lines = text.lower().splitlines()
    sections = {}
    current_section = "unknown_introduction" # Default section for text before any recognized header
    buffer = []

    for line in lines:
        found_header = False
        for h in headers:
            # Check if the line contains a header, allowing for some flexibility (e.g., "History of Present Illness:")
            if h in line and len(line) < 100: # Heuristic: header lines are usually short
                if buffer:
                    sections[current_section.strip()] = "\n".join(buffer).strip()
                    buffer = []
                current_section = line.strip()
                found_header = True
                break
        if not found_header:
            buffer.append(line)

    # Add any remaining text to the last section or 'unknown_introduction'
    if buffer:
        sections[current_section.strip()] = "\n".join(buffer).strip()
    return sections


def categorize_section_type(text):
    """
    Uses Gemini to categorize the type of clinical documentation section.
    Provides a document type and a confidence score.
    """
    snippet = text[:1500] # Use a larger snippet for better context
    prompt = f"""
    You are an expert in clinical documentation. Analyze the following text snippet from a patient's chart.
    What type of clinical documentation is this?
    Choose one from the most common types: History and Physical (H&P), Progress Note, Operative Report, Discharge Summary, Consultation Report, Emergency Department Record, Diagnostic Test Results, Physician Orders, Anesthesiology Record, Medication Administration Record, Vital Signs Flowsheet, Problem List, Ventilator Flowsheet, Wound Care Notes, or Other/Unknown.
    Provide a confidence percentage for your categorization.

    --- Text Snippet ---
    {snippet}
    --- End Snippet ---
    Respond in JSON format like this: {{"Document Type": "str", "Confidence": int}}
    """
    try:
        response = model.generate_content(prompt)
        # Attempt to parse JSON, handle cases where Gemini might return extra text
        json_str = response.text.strip().replace("```json\n", "").replace("\n```", "")
        return json.loads(json_str)
    except Exception as e:
        st.warning(f"Error categorizing section type: {e}. Gemini response: {response.text}")
        return {"Document Type": "Unknown", "Confidence": 0}


def generate_gemini_analysis(text, retries=3):
    """
    Uses Gemini to perform a detailed CDI analysis on the provided clinical text.
    Evaluates against 5 CDI criteria, provides scores, explanations, proposed actions,
    and draft queries.
    """
    # Detailed definitions of CDI criteria, derived from the uploaded RTF documents,
    # embedded directly into the prompt to guide Gemini's reasoning.
    cdi_criteria_definitions = """
    Here are the definitions for the 5 CDI-relevant criteria you must assess, drawing from the principles of high-quality clinical documentation:

    1.  **Clinical Reliability (Score 0-10):**
        * **Definition:** The content of the record is trustworthy, safe, and internally consistent, supporting the treatments provided. Documentation should logically align with the patient's condition and the interventions performed.
        * **Example of Deficiency:** Treatment provided without documentation of the condition being treated (e.g., Lasix given, but no CHF documented; KCI administered, no hypokalemia documented). Or, a diagnosis that does not appear reliable based on the treatment (e.g., blood transfusion for a bleeding gastric ulcer without acute blood loss anemia).
        * **Goal:** Ensure diagnoses and treatments are mutually supportive and clinically sound.

    2.  **Clinical Precision (Score 0-10):**
        * **Definition:** The record is accurate, exact, and strictly defined. Documentation should include specific details, etiology, and highest level of specificity where clinically appropriate. Avoid vague or general terms.
        * **Example of Deficiency:** No specific diagnosis documented when a more specific diagnosis appears to be supported (e.g., "anemia" vs. "acute or chronic blood loss anemia"; "pneumonia" vs. "aspiration pneumonia").
        * **Goal:** Drive documentation to the highest level of detail and specificity for accurate coding and reflection of patient acuity.

    3.  **Documentation Completeness (Score 0-10):**
        * **Definition:** The record has the maximum content and is thorough. All concerns are fully addressed, including appropriate authentication (signature, date). Abnormal test results should have documented clinical significance or be noted as clinically insignificant. Reasons for tests/treatments should be documented.
        * **Example of Deficiency:** Abnormal test results without documentation for clinical significance (e.g., low sodium, magnesium, potassium without corresponding diagnoses or notation of clinical insignificance). Lack of documented reason for ordered tests or treatments.
        * **Goal:** Ensure all relevant clinical information is present, authenticated, and fully explained.

    4.  **Clinical Consistency & Clarity (Score 0-10):**
        * **Definition:** The record is not contradictory, vague, or ambiguous. Documentation should be coherent across different entries and authors. If conflicting information exists, it should be resolved or clarified. Symptoms should ideally have identified etiologies; if not, "etiology undetermined" should be documented.
        * **Example of Deficiency:** Disagreement between two or more treating physicians without obvious resolution (e.g., primary care documents TIA, neurologist documents CVA, and attending physician doesn't clarify). Vague or ambiguous documentation, especially for symptom principal diagnoses (e.g., "chest pain" without further insight; "syncope" without etiology).
        * **Goal:** Achieve a unified, unambiguous clinical picture that accurately reflects the patient's condition and care.

    5.  **Timeliness (Flag Only - "Present", "Absent", or "N/A"):**
        * **Definition:** Documentation is prepared by the provider at the point it is needed for patient care, adhering to facility guidelines. This includes timely progress notes, discharge summaries, and documentation of Present on Admission (POA) indicators.
        * **Flagging:** You will only flag if there are *clear indicators* of timeliness issues (e.g., explicit mentions of late documentation, missing dates where expected). If no such indicators are present, assume "Absent" or "N/A" if the document type doesn't typically have strict timeliness requirements visible in the text.
        * **Goal:** Ensure documentation is current and available when needed for care, coding, and compliance.
    """

    prompt = f"""
    You are a Senior Clinical Documentation Integrity (CDI) Specialist. Your task is to systematically assess the following clinical document section using the provided theory of high-quality documentation.
    Apply your expertise to validate the documentation and identify any disturbances or areas for improvement.

    {cdi_criteria_definitions}

    For each of the first four criteria (Clinical Reliability, Clinical Precision, Documentation Completeness, Clinical Consistency & Clarity):
    - Assign a score (0-10), where 10 is perfect and 0 is completely deficient.
    - Explain "Why" this score was given, referencing specific issues or strengths in the document.
    - Suggest a "Proposed CDI Action" (e.g., "Educate provider on specificity," "Query for clinical validation," "Review for conflicting documentation").
    - If the score is less than 10, draft a clear and compliant "Draft Query" that a CDI specialist would send to the physician to clarify or improve the documentation. The query should be open-ended and non-leading.

    For the "Timeliness" criterion:
    - Flag it as "Present" if there are clear indications of timeliness issues (e.g., explicit mentions of late entries, missing dates where expected).
    - Flag it as "Absent" if no such issues are evident from the text.
    - Flag it as "N/A" if the document type or content doesn't provide enough information to assess timeliness.

    Return your output in structured JSON like this. Ensure all fields are present, even if empty or "N/A".

    {{
      "Scorecard": {{
        "Clinical Reliability": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Clinical Precision": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Documentation Completeness": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}},
        "Clinical Consistency & Clarity": {{"Score": int, "Why": "str", "Proposed CDI Action": "str", "Draft Query": "str"}}
      }},
      "Timeliness Flag": "Present" | "Absent" | "N/A",
      "Final Summary": "Overall assessment of documentation quality for this section, highlighting key strengths and areas for CDI focus."
    }}

    --- Document Section Start ---
    {text}
    --- Document Section End ---
    """
    for attempt in range(retries):
        try:
            response = model.generate_content(prompt)
            # Clean the response to ensure it's valid JSON
            json_str = response.text.strip().replace("```json\n", "").replace("\n```", "")
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            st.warning(f"Attempt {attempt+1}: JSON decoding error: {e}. Retrying...")
            time.sleep(2) # Wait before retrying
        except Exception as e:
            st.warning(f"Attempt {attempt+1}: Gemini generation error: {e}. Retrying...")
            time.sleep(2) # Wait before retrying
    return {"error": "Gemini failed after retries or returned invalid JSON.", "raw_response": response.text if 'response' in locals() else "No response"}


def generate_excel(results):
    """
    Generates a Pandas DataFrame from the analysis results, suitable for Excel export.
    Includes all CDI criteria, scores, explanations, actions, and queries.
    """
    df_rows = []
    for section, result in results.items():
        scorecard = result.get("Scorecard", {})
        doc_type = result.get("Doc Type", "")
        confidence = result.get("Confidence", "")
        timeliness_flag = result.get("Timeliness Flag", "N/A")
        final_summary = result.get("Final Summary", "")

        # Add a row for the overall section summary and timeliness flag
        df_rows.append({
            "Section": section.title(),
            "AI Doc Type": doc_type,
            "Confidence %": confidence,
            "Criterion": "Overall Section Summary",
            "Score": "-", # No score for summary row
            "Why": final_summary,
            "Proposed CDI Action": "",
            "Draft Query": "",
            "Timeliness Flag": timeliness_flag # Display timeliness here
        })

        # Add rows for each detailed scorecard criterion
        for criterion, detail in scorecard.items():
            score = detail.get("Score", "")
            # Only include Draft Query if score is less than 10
            draft_query = detail.get("Draft Query", "") if int(str(score).replace('-', '0')) < 10 else ""
            df_rows.append({
                "Section": section.title(),
                "AI Doc Type": doc_type,
                "Confidence %": confidence,
                "Criterion": criterion,
                "Score": score,
                "Why": detail.get("Why", ""),
                "Proposed CDI Action": detail.get("Proposed CDI Action", ""),
                "Draft Query": draft_query,
                "Timeliness Flag": timeliness_flag # Repeat for consistency in rows
            })
    return pd.DataFrame(df_rows)

def highlight_score(val):
    """
    Applies conditional formatting to the 'Score' column in the DataFrame.
    Green for good scores, yellow for moderate, red for low.
    """
    try:
        score = int(val)
        if score >= 9:
            return 'background-color: #d9ead3'  # Light Green
        elif score >= 7:
            return 'background-color: #fff2cc'  # Light Yellow
        else:
            return 'background-color: #fce5cd'  # Light Orange
    except ValueError: # Handle non-integer values like '-'
        return ''

# --- Streamlit UI ---
st.set_page_config(page_title="Inpatient CDI Analyzer", layout="wide")
st.title("🩺 Inpatient CDI Analyzer: AI-Powered Clinical Documentation Review")
st.markdown("Upload a full inpatient PDF or Word document to get a detailed CDI analysis, including scores, proposed actions, and draft queries.")

uploaded_file = st.file_uploader("📤 Upload Full Inpatient Chart (.pdf or .docx)", type=["pdf", "docx"])

if uploaded_file:
    with st.spinner("Processing document..."):
        if uploaded_file.name.endswith(".pdf"):
            full_text = extract_text_from_pdf(uploaded_file)
        else:
            full_text = extract_text_from_docx(uploaded_file)

    st.success("✅ Document Uploaded and Processed!")

    sections = split_sections(full_text)
    st.subheader("📑 Detected Sections")
    st.write(f"Found {len(sections)} sections in the document.")

    all_results = {}

    # Button to analyze all sections
    if st.button("🧠 Analyze All Sections (This may take a few minutes for large documents)"):
        progress_bar = st.progress(0)
        total_sections = len(sections)
        for i, (section_name, section_text) in enumerate(sections.items()):
            st.info(f"Analyzing section: **{section_name.title()}** ({i+1}/{total_sections})...")
            with st.spinner(f"Analyzing: {section_name.title()}..."):
                doc_info = categorize_section_type(section_text)
                analysis_result = generate_gemini_analysis(section_text)

                if 'error' not in analysis_result and analysis_result.get("Scorecard"):
                    analysis_result["Doc Type"] = doc_info.get("Document Type")
                    analysis_result["Confidence"] = doc_info.get("Confidence")
                    all_results[section_name] = analysis_result
                    st.success(f"✅ Analysis complete for **{section_name.title()}**.")
                else:
                    st.error(f"❌ Failed to analyze section: **{section_name.title()}**. Error: {analysis_result.get('error', 'Unknown')}")
            progress_bar.progress((i + 1) / total_sections)
        st.success("🎉 All sections analyzed!")

        # Display results in a DataFrame
        if all_results:
            st.subheader("📊 Comprehensive CDI Analysis Results")
            df = generate_excel(all_results)

            if 'Score' in df.columns:
                # Apply styling only to rows that are actual criterion scores, not summary rows
                styled_df = df.style.applymap(highlight_score, subset=['Score'])
                st.dataframe(styled_df, use_container_width=True)
            else:
                st.warning("⚠️ No valid 'Score' column found in results. Showing raw output.")
                st.dataframe(df, use_container_width=True)

            # Download buttons
            col1, col2 = st.columns(2)
            with col1:
                with BytesIO() as buffer:
                    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
                        df.to_excel(writer, index=False)
                    st.download_button("📥 Download Excel Report", buffer.getvalue(), file_name="CDI_Analysis_Report.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            with col2:
                txt_summary = ""
                for sec, res in all_results.items():
                    txt_summary += f"\n\n=== {sec.title()} (AI Doc Type: {res.get('Doc Type', '')}, Confidence: {res.get('Confidence', '')}%) ===\n"
                    txt_summary += f"Overall Summary: {res.get('Final Summary', 'N/A')}\n"
                    txt_summary += f"Timeliness Flag: {res.get('Timeliness Flag', 'N/A')}\n"
                    if 'error' in res:
                        txt_summary += "Error during analysis.\n"
                    else:
                        for k, v in res.get("Scorecard", {}).items():
                            score = v.get('Score', '-')
                            why = v.get('Why', '')
                            action = v.get('Proposed CDI Action', '')
                            query = v.get('Draft Query', '')

                            txt_summary += f"\n--- {k} (Score: {score}/10) ---\n"
                            txt_summary += f"Why: {why}\n"
                            txt_summary += f"Proposed CDI Action: {action}\n"
                            if query: # Only add query if it's not empty
                                txt_summary += f"→ Draft Query: {query}\n"
                st.download_button("📄 Download TXT Report", txt_summary, file_name="CDI_Queries_Report.txt", mime="text/plain")
        else:
            st.info("No analysis results to display yet. Upload a document and click 'Analyze All Sections'.")

    st.markdown("---")
    st.subheader("🔍 Analyze Individual Sections")
    st.info("You can also expand a section below and analyze it individually.")

    # Display individual sections with an analyze button for each
    for i, (section_name, section_text) in enumerate(sections.items()):
        with st.expander(f"📄 Section {i+1}: {section_name.title()}"):
            st.text_area(f"Content of '{section_name.title()}' (first 2000 chars)", section_text[:2000], height=200, key=f"text_area_{i}")
            if st.button(f"🔍 Analyze '{section_name.title()}' Individually", key=f"analyze_single_{i}"):
                with st.spinner(f"Analyzing '{section_name.title()}' with Gemini..."):
                    doc_info = categorize_section_type(section_text)
                    result = generate_gemini_analysis(section_text)
                    result["Doc Type"] = doc_info.get("Document Type")
                    result["Confidence"] = doc_info.get("Confidence")

                    if 'error' not in result:
                        st.success(f"✅ Analysis for '{section_name.title()}' complete!")
                        st.json(result) # Display raw JSON for single section analysis
                    else:
                        st.error(f"❌ Failed to analyze '{section_name.title()}'. Error: {result.get('error', 'Unknown')}")
                        if 'raw_response' in result:
                            st.code(result['raw_response'], language='json') # Show raw response for debugging
